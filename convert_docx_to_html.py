#!/usr/bin/env python3
"""
Word Document to HTML Converter for yanivgil.co.il
Converts pillar content Word documents to HTML pages with RTL Hebrew formatting.
INCLUDES ALL CONTENT - No shortening or editing of source content.
"""

import os
import re
import sys
import io

# Fix Unicode output on Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# Configuration
INPUT_FOLDER = r'c:\Users\user\Desktop\yaniv\New_content'
OUTPUT_FOLDER = r'c:\Users\user\Desktop\yaniv'
TEMPLATE_URL = 'https://www.yanivgil.co.il'

# Skip duplicate files
SKIP_FILES = ['(1).docx', 'SEO Content Strategy']


def get_html_template(title, description, keyword, url, content, breadcrumb_pillar='', breadcrumb_pillar_url='', sub_pages=None):
    """Generate full HTML page from template - matches index.html design exactly."""

    # Clean URL for canonical
    clean_url = url.strip('/').replace('.html', '') + '.html'
    full_url = f'{TEMPLATE_URL}/{clean_url}'

    # Generate breadcrumb schema
    breadcrumb_items = [
        '{"@type": "ListItem", "position": 1, "name": "בית", "item": "https://www.yanivgil.co.il/"}',
        '{"@type": "ListItem", "position": 2, "name": "דיני משפחה", "item": "https://www.yanivgil.co.il/family-law.html"}',
        '{"@type": "ListItem", "position": 3, "name": "המדריך המלא לגירושין", "item": "https://www.yanivgil.co.il/divorce-guide.html"}'
    ]

    if breadcrumb_pillar:
        breadcrumb_items.append(f'{{"@type": "ListItem", "position": 4, "name": "{breadcrumb_pillar}", "item": "{TEMPLATE_URL}/{breadcrumb_pillar_url}"}}')
        breadcrumb_items.append(f'{{"@type": "ListItem", "position": 5, "name": "{title}"}}')
    else:
        breadcrumb_items.append(f'{{"@type": "ListItem", "position": 4, "name": "{title}"}}')

    breadcrumb_schema = ',\n    '.join(breadcrumb_items)

    # Build breadcrumb HTML
    breadcrumb_html = '''<nav class="breadcrumbs" aria-label="Breadcrumb">
            <ol class="flex flex-wrap items-center gap-2 text-sm text-gray-600">
                <li><a href="index.html" class="hover:text-[#c9a55c]">בית</a></li>
                <li><span class="mx-2">/</span></li>
                <li><a href="family-law.html" class="hover:text-[#c9a55c]">דיני משפחה</a></li>
                <li><span class="mx-2">/</span></li>
                <li><a href="divorce-guide.html" class="hover:text-[#c9a55c]">המדריך המלא לגירושין</a></li>'''

    if breadcrumb_pillar:
        breadcrumb_html += f'''
                <li><span class="mx-2">/</span></li>
                <li><a href="{breadcrumb_pillar_url}" class="hover:text-[#c9a55c]">{breadcrumb_pillar}</a></li>'''

    breadcrumb_html += f'''
                <li><span class="mx-2">/</span></li>
                <li class="text-[#c9a55c] font-medium" aria-current="page">{title}</li>
            </ol>
        </nav>'''

    # Build sub-pages navigation for pillar pages
    sub_pages_html = ''
    if sub_pages and len(sub_pages) > 0:
        sub_pages_html = '''
            <div class="bg-gray-50 rounded-xl p-6 mt-8 mb-6">
                <h3 class="text-xl font-bold text-[#1a1a1a] mb-4">מאמרים נוספים בנושא</h3>
                <div class="grid grid-cols-1 md:grid-cols-2 gap-3">'''
        for sp in sub_pages:
            sub_pages_html += f'''
                    <a href="{sp['url']}" class="flex items-center gap-2 p-3 bg-white rounded-lg border border-gray-200 hover:border-[#c9a55c] hover:shadow-md transition-all">
                        <svg class="w-4 h-4 text-[#c9a55c] flex-shrink-0" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9 5l7 7-7 7"/></svg>
                        <span class="text-gray-700 hover:text-[#c9a55c]">{sp['title']}</span>
                    </a>'''
        sub_pages_html += '''
                </div>
            </div>'''

    return f'''<!DOCTYPE html>
<html lang="he" dir="rtl">
<head>
    <!-- Google tag (gtag.js) -->
    <script async src="https://www.googletagmanager.com/gtag/js?id=G-1ES9G9LMG6"></script>
    <script>
      window.dataLayer = window.dataLayer || [];
      function gtag(){{dataLayer.push(arguments);}}
      gtag('js', new Date());
      gtag('config', 'G-1ES9G9LMG6');
    </script>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title} | עו"ד יניב גיל - תל אביב</title>
    <meta name="description" content="{description}">
    <meta name="keywords" content="{keyword}, עורך דין דיני משפחה, עורך דין גירושין תל אביב">
    <meta property="og:title" content="{title} | עו״ד יניב גיל">
    <meta property="og:description" content="{description}">
    <meta property="og:type" content="article">
    <meta property="og:url" content="{full_url}">
    <meta property="og:image" content="https://www.yanivgil.co.il/images/new-logo.png">
    <meta property="og:locale" content="he_IL">
    <meta property="og:site_name" content="משרד עו״ד יניב גיל ושות'">
    <link rel="canonical" href="{full_url}">
    <!-- Article Schema -->
    <script type="application/ld+json">
    {{
      "@context": "https://schema.org",
      "@type": "Article",
      "headline": "{title}",
      "author": {{"@type": "Person", "name": "עו״ד יניב גיל", "url": "https://www.yanivgil.co.il/attorneys.html"}},
      "publisher": {{"@type": "Organization", "name": "משרד עו״ד יניב גיל ושות׳", "logo": {{"@type": "ImageObject", "url": "https://www.yanivgil.co.il/images/new-logo.png"}}}},
      "datePublished": "{datetime.now().strftime('%Y-%m-%d')}",
      "dateModified": "{datetime.now().strftime('%Y-%m-%d')}",
      "mainEntityOfPage": "{full_url}",
      "description": "{description}"
    }}
    </script>
    <!-- BreadcrumbList Schema -->
    <script type="application/ld+json">
    {{
      "@context": "https://schema.org",
      "@type": "BreadcrumbList",
      "itemListElement": [
        {breadcrumb_schema}
      ]
    }}
    </script>
    <link rel="icon" type="image/x-icon" href="favicon.ico">
    <link rel="stylesheet" href="css/tailwind.css">
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link rel="dns-prefetch" href="https://www.google-analytics.com">
    <link rel="dns-prefetch" href="https://www.googletagmanager.com">
    <link href="https://fonts.googleapis.com/css2?family=Heebo:wght@300;400;500;700&display=swap" rel="stylesheet">
    <link rel="stylesheet" href="css/mobile-menu.css">
    <link rel="stylesheet" href="css/header-responsive.css">
    <link rel="stylesheet" href="css/breadcrumbs.css">
    <link rel="stylesheet" href="css/back-to-top.css">
    <link rel="stylesheet" href="css/progress-bar.css">
    <link rel="stylesheet" href="css/floating-buttons.css">
    <link rel="stylesheet" href="css/print.css" media="print">
    <link rel="stylesheet" href="css/mobile-cta-bar.css">
    <style>
        html {{ overflow-x: hidden !important; max-width: 100% !important; }}
        body {{ overflow-x: hidden !important; max-width: 100% !important; font-family: 'Heebo', sans-serif; direction: rtl; }}
        * {{ box-sizing: border-box !important; max-width: 100vw; }}
        .mobile-menu-overlay {{ position: fixed; top: 0; left: 0; width: 100%; height: 100%; background-color: #ffffff; z-index: 9999; transform: translateX(100%); transition: transform 0.3s ease-in-out; overflow-y: auto; display: flex; flex-direction: column; }}
        .mobile-menu-overlay.active {{ transform: translateX(0); }}
        .mobile-menu-overlay > .flex.justify-between {{ background: #ffffff; border-bottom: 1px solid #e5e7eb; }}
        .mobile-menu-overlay nav {{ background: #ffffff; flex: 1; }}
        .mobile-menu-overlay nav ul {{ list-style: none; margin: 0; padding: 0; }}
        .mobile-menu-overlay nav ul li a.mobile-menu-link {{ display: block; padding: 1rem; background: #1a1a1a !important; color: #c9a55c !important; text-decoration: none; border-radius: 5px; margin-bottom: 0.5rem; transition: all 0.3s ease; }}
        .mobile-menu-overlay nav ul li a.mobile-menu-link:hover {{ background: #c9a55c !important; color: #1a1a1a !important; }}
        .tip-box {{ border-right: 4px solid #ca8a04; background: #fefce8; padding: 1rem 1.5rem; border-radius: 0.5rem; margin: 1.5rem 0; }}
        .info-box {{ border-right: 4px solid #2563eb; background: #eff6ff; padding: 1rem 1.5rem; border-radius: 0.5rem; margin: 1.5rem 0; }}
        .warning-box {{ border-right: 4px solid #dc2626; background: #fef2f2; padding: 1rem 1.5rem; border-radius: 0.5rem; margin: 1.5rem 0; }}
        .cta-box {{ background: linear-gradient(135deg, #1a1a2e, #16213e); color: white; padding: 2rem; border-radius: 1rem; text-align: center; margin: 2rem 0; }}
        .cta-box .cta-btn-call {{ background: #c9a55c !important; color: #1a1a1a !important; text-decoration: none !important; }}
        .cta-box .cta-btn-call:hover {{ background: #ffffff !important; }}
        .cta-box .cta-btn-whatsapp {{ background: #16a34a !important; color: #ffffff !important; text-decoration: none !important; }}
        .cta-box .cta-btn-whatsapp:hover {{ background: #15803d !important; }}
        .author-box {{ background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 0.75rem; padding: 1.5rem; margin: 2rem 0; }}
        .content-section {{ padding: 1.5rem 0; }}
        .content-section h2 {{ color: #1a1a1a; font-size: 1.5rem; font-weight: 700; margin-bottom: 1rem; padding-bottom: 0.5rem; border-bottom: 2px solid #c9a55c; }}
        .content-section h3 {{ color: #374151; font-size: 1.25rem; font-weight: 600; margin: 1.5rem 0 0.75rem; }}
        .content-section p {{ color: #4b5563; line-height: 1.8; margin-bottom: 1rem; }}
        .content-section ul {{ list-style: none; padding-right: 0; margin: 1rem 0; }}
        .content-section ul li {{ position: relative; padding-right: 1.5rem; margin-bottom: 0.75rem; color: #4b5563; line-height: 1.7; }}
        .content-section ul li::before {{ content: "•"; position: absolute; right: 0; color: #c9a55c; font-weight: bold; }}
        .content-section ol {{ padding-right: 1.5rem; margin: 1rem 0; }}
        .content-section ol li {{ margin-bottom: 0.75rem; color: #4b5563; line-height: 1.7; }}
        .content-table {{ width: 100%; border-collapse: collapse; margin: 1.5rem auto; background: white; border-radius: 0.5rem; overflow: hidden; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }}
        .content-table th {{ background: #1a1a1a; color: #c9a55c; padding: 1rem; text-align: right; font-weight: 600; }}
        .content-table td {{ padding: 0.75rem 1rem; border-bottom: 1px solid #e5e7eb; text-align: right; }}
        .content-table tr:hover {{ background: #f9fafb; }}
        .back-link {{ display: inline-flex; align-items: center; gap: 0.5rem; color: #c9a55c; font-weight: 500; margin-bottom: 1rem; text-decoration: none; }}
        .back-link:hover {{ text-decoration: underline; }}
        @media (max-width: 768px) {{ .whatsapp-float, .call-float, .accessibility-btn, .reading-progress {{ display: none !important; }} }}
        footer, footer * {{ color: #d1d5db !important; }}
        footer h3, footer h4 {{ color: #ffffff !important; }}
        footer a:hover {{ color: #c9a55c !important; }}
    </style>
</head>
<body class="bg-gray-50 text-gray-800">
    <!-- Reading Progress Bar -->
    <div class="reading-progress" role="progressbar" aria-label="התקדמות קריאה"></div>

    <!-- Header - matches index.html exactly -->
    <header class="site-header fixed top-0 left-0 w-full z-50 bg-white transition-all duration-300 ease-in-out py-6 shadow-md" role="banner" id="site-header">
        <div class="container mx-auto px-4 flex justify-between items-center transition-all duration-300 ease-in-out">
            <div style="display: flex; align-items: center; gap: 15px;">
                <a href="index.html" title="דף הבית"><img src="images/new-logo.png" alt="לוגו משרד עו״ד יניב גיל" class="header-logo" style="margin-left: 10px; display: block;" width="60" height="60" loading="eager"></a>
                <a href="index.html" title="דף הבית" class="flex items-center gap-2 text-gray-800 transition-all duration-300 ease-in-out"><img src="images/Yaniv-Gil-Law-Office-Notary_text.avif" alt="משרד עו״ד יניב גיל ושות׳" class="header-text-logo" width="220" height="32" loading="eager"></a>
            </div>
            <nav class="hidden md:flex flex-1 justify-center" role="navigation" aria-label="תפריט ראשי">
                <ul class="flex gap-6 list-none m-0 p-0 text-gray-600 font-medium">
                    <li><a href="index.html" title="דף הבית" class="inline-flex items-center justify-center rounded-full border border-[#c9a55c] bg-[#1a1a1a] px-4 py-2 text-sm font-medium text-[#c9a55c] shadow-sm transition-all duration-300 hover:bg-[#c9a55c] hover:text-[#1a1a1a]">בית</a></li>
                    <li><a href="index.html#about" title="אודות המשרד" class="inline-flex items-center justify-center rounded-full border border-[#c9a55c] bg-[#1a1a1a] px-4 py-2 text-sm font-medium text-[#c9a55c] shadow-sm transition-all duration-300 hover:bg-[#c9a55c] hover:text-[#1a1a1a]">אודות המשרד</a></li>
                    <li><a href="attorneys.html" title="עורכי דין" class="inline-flex items-center justify-center rounded-full border border-[#c9a55c] bg-[#1a1a1a] px-4 py-2 text-sm font-medium text-[#c9a55c] shadow-sm transition-all duration-300 hover:bg-[#c9a55c] hover:text-[#1a1a1a]">עורכי דין</a></li>
                    <li class="relative group">
                        <button onclick="toggleNavDropdown('practice-dropdown')" class="inline-flex items-center justify-center rounded-full border border-[#c9a55c] bg-[#1a1a1a] px-4 py-2 text-sm font-medium text-[#c9a55c] shadow-sm transition-all duration-300 hover:bg-[#c9a55c] hover:text-[#1a1a1a]">תחומי התמחות <span class="mr-2 text-xs">▼</span></button>
                        <div id="practice-dropdown" class="hidden absolute top-full mt-2 right-0 w-64 bg-white rounded-lg shadow-xl border border-gray-200 z-50 py-2" style="max-height: calc(100vh - 120px); overflow-y: auto;">
                            <!-- דיני משפחה - highlighted -->
                            <div class="border-b border-gray-200 bg-amber-50">
                                <button onclick="toggleCategoryDropdown('category-family', event)" class="w-full px-4 py-3 flex items-center justify-between text-right hover:bg-amber-100" style="background: linear-gradient(90deg, #fef3c7, #fff);"><h4 class="font-bold text-[#b8941f] text-sm">דיני משפחה וגירושין</h4><span class="category-icon text-lg text-[#c9a55c] font-bold">▼</span></button>
                                <div id="category-family" class="hidden">
                                    <a href="divorce-guide.html" class="block px-4 py-2 pr-6 text-sm text-[#b8941f] font-bold hover:bg-amber-100 text-right border-b border-gray-100">★ המדריך המלא לגירושין</a>
                                    <div class="pillar-item border-b border-gray-100"><div class="flex items-center justify-between pr-6 hover:bg-gray-50"><a href="heskem-gerushin.html" class="flex-1 py-2 text-sm text-gray-700 hover:text-[#b8941f] font-medium">הסכם גירושין</a><button onclick="toggleSubCategory('sub-heskem-gerushin', event)" class="px-3 py-2 text-lg text-[#c9a55c] font-bold hover:text-[#b8941f]">▼</button></div><div id="sub-heskem-gerushin" class="hidden bg-gray-50 pr-10 py-1"><a href="mahu-heskem-gerushin.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">מהו הסכם גירושין</a><a href="shlvei-heskem-gerushin.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">שלבי הסכם</a><a href="seifim-hova-heskem.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">סעיפי חובה</a></div></div>
                                    <div class="pillar-item border-b border-gray-100"><div class="flex items-center justify-between pr-6 hover:bg-gray-50"><a href="gerushin-beit-din-rabani.html" class="flex-1 py-2 text-sm text-gray-700 hover:text-[#b8941f] font-medium">בית דין רבני</a><button onclick="toggleSubCategory('sub-beit-din', event)" class="px-3 py-2 text-lg text-[#c9a55c] font-bold hover:text-[#b8941f]">▼</button></div><div id="sub-beit-din" class="hidden bg-gray-50 pr-10 py-1"><a href="haget-mah-ze.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">הגט - מה זה</a><a href="smchut-beit-din-rabani.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">סמכות בית דין</a></div></div>
                                    <div class="pillar-item border-b border-gray-100"><div class="flex items-center justify-between pr-6 hover:bg-gray-50"><a href="mishmoret-yeladim.html" class="flex-1 py-2 text-sm text-gray-700 hover:text-[#b8941f] font-medium">משמורת ילדים</a><button onclick="toggleSubCategory('sub-mishmoret', event)" class="px-3 py-2 text-lg text-[#c9a55c] font-bold hover:text-[#b8941f]">▼</button></div><div id="sub-mishmoret" class="hidden bg-gray-50 pr-10 py-1"><a href="mishmoret-blaudit-vsmeshuteft.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">בלעדית vs משותפת</a><a href="achrayut-horit.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">אחריות הורית</a></div></div>
                                    <div class="pillar-item border-b border-gray-100"><div class="flex items-center justify-between pr-6 hover:bg-gray-50"><a href="mezonot-yeladim.html" class="flex-1 py-2 text-sm text-gray-700 hover:text-[#b8941f] font-medium">מזונות ילדים</a><button onclick="toggleSubCategory('sub-mezonot-yeladim', event)" class="px-3 py-2 text-lg text-[#c9a55c] font-bold hover:text-[#b8941f]">▼</button></div><div id="sub-mezonot-yeladim" class="hidden bg-gray-50 pr-10 py-1"><a href="chishov-mezonot-yeladim.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">חישוב מזונות</a><a href="hafchatat-mezonot-yeladim.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">הפחתת מזונות</a></div></div>
                                    <div class="pillar-item border-b border-gray-100"><div class="flex items-center justify-between pr-6 hover:bg-gray-50"><a href="mezonot-isha.html" class="flex-1 py-2 text-sm text-gray-700 hover:text-[#b8941f] font-medium">מזונות אישה</a><button onclick="toggleSubCategory('sub-mezonot-isha', event)" class="px-3 py-2 text-lg text-[#c9a55c] font-bold hover:text-[#b8941f]">▼</button></div><div id="sub-mezonot-isha" class="hidden bg-gray-50 pr-10 py-1"><a href="zkaut-mezonot-isha.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">זכאות למזונות</a><a href="chishov-mezonot-isha.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">חישוב מזונות</a></div></div>
                                    <div class="pillar-item border-b border-gray-100"><div class="flex items-center justify-between pr-6 hover:bg-gray-50"><a href="chalokat-rechush-gerushin.html" class="flex-1 py-2 text-sm text-gray-700 hover:text-[#b8941f] font-medium">חלוקת רכוש</a><button onclick="toggleSubCategory('sub-rechush', event)" class="px-3 py-2 text-lg text-[#c9a55c] font-bold hover:text-[#b8941f]">▼</button></div><div id="sub-rechush" class="hidden bg-gray-50 pr-10 py-1"><a href="izun-mashabim-gerushin.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">איזון משאבים</a><a href="chalokat-dira-gerushin.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">חלוקת דירה</a></div></div>
                                    <div class="pillar-item border-b border-gray-100"><div class="flex items-center justify-between pr-6 hover:bg-gray-50"><a href="heskem-mamon.html" class="flex-1 py-2 text-sm text-gray-700 hover:text-[#b8941f] font-medium">הסכם ממון</a><button onclick="toggleSubCategory('sub-mamon', event)" class="px-3 py-2 text-lg text-[#c9a55c] font-bold hover:text-[#b8941f]">▼</button></div><div id="sub-mamon" class="hidden bg-gray-50 pr-10 py-1"><a href="heskem-mamon-lifney-nisuin.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">לפני נישואין</a><a href="hagana-al-esek.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">הגנה על עסק</a></div></div>
                                    <div class="pillar-item border-b border-gray-100"><div class="flex items-center justify-between pr-6 hover:bg-gray-50"><a href="sarvanut-get-agunot.html" class="flex-1 py-2 text-sm text-gray-700 hover:text-[#b8941f] font-medium">סרבנות גט</a><button onclick="toggleSubCategory('sub-sarvanut', event)" class="px-3 py-2 text-lg text-[#c9a55c] font-bold hover:text-[#b8941f]">▼</button></div><div id="sub-sarvanut" class="hidden bg-gray-50 pr-10 py-1"><a href="klim-mishpatiyim-sarvan-get.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">כלים משפטיים</a><a href="agunut-ma-laasot.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">עגינות - מה לעשות</a></div></div>
                                    <div class="pillar-item border-b border-gray-100"><div class="flex items-center justify-between pr-6 hover:bg-gray-50"><a href="yaduyim-batzibur.html" class="flex-1 py-2 text-sm text-gray-700 hover:text-[#b8941f] font-medium">ידועים בציבור</a><button onclick="toggleSubCategory('sub-yaduyim', event)" class="px-3 py-2 text-lg text-[#c9a55c] font-bold hover:text-[#b8941f]">▼</button></div><div id="sub-yaduyim" class="hidden bg-gray-50 pr-10 py-1"><a href="zkuyot-yaduyim-batzibur.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">זכויות</a><a href="heskem-yaduyim-batzibur.html" class="block py-1.5 px-4 text-xs text-gray-600 hover:text-[#b8941f]">הסכם חיים משותפים</a></div></div>
                                    <a href="divorce.html" class="block px-4 py-2 pr-6 text-sm text-gray-600 hover:bg-gray-50 hover:text-[#b8941f] text-right">גירושין</a>
                                    <a href="guardianship.html" class="block px-4 py-2 pr-6 text-sm text-gray-600 hover:bg-gray-50 hover:text-[#b8941f] text-right">אפוטרופסות</a>
                                    <a href="mediation.html" class="block px-4 py-2 pr-6 text-sm text-gray-600 hover:bg-gray-50 hover:text-[#b8941f] text-right">גישור משפחתי</a>
                                </div>
                            </div>
                            <!-- ירושות וצוואות -->
                            <div class="border-b border-gray-200 mt-2">
                                <button onclick="toggleCategoryDropdown('category-inheritance', event)" class="w-full px-4 py-2 flex items-center justify-between text-right hover:bg-gray-50"><h4 class="font-bold text-[#1a1a1a] text-sm">ירושות וצוואות</h4><span class="category-icon text-xs text-gray-500">▼</span></button>
                                <div id="category-inheritance" class="hidden">
                                    <a href="will-writing.html" class="block px-4 py-2 pr-8 text-sm text-gray-700 hover:bg-[#c9a55c] hover:text-[#1a1a1a] text-right">עריכת צוואה</a>
                                    <a href="inheritance-order.html" class="block px-4 py-2 pr-8 text-sm text-gray-700 hover:bg-[#c9a55c] hover:text-[#1a1a1a] text-right">צו ירושה</a>
                                    <a href="will-probate.html" class="block px-4 py-2 pr-8 text-sm text-gray-700 hover:bg-[#c9a55c] hover:text-[#1a1a1a] text-right">צו קיום צוואה</a>
                                    <a href="estate-management.html" class="block px-4 py-2 pr-8 text-sm text-gray-700 hover:bg-[#c9a55c] hover:text-[#1a1a1a] text-right">ניהול עיזבון</a>
                                </div>
                            </div>
                            <!-- חדלות פירעון -->
                            <div class="border-b border-gray-200 mt-2">
                                <button onclick="toggleCategoryDropdown('category-bankruptcy', event)" class="w-full px-4 py-2 flex items-center justify-between text-right hover:bg-gray-50"><h4 class="font-bold text-[#1a1a1a] text-sm">חדלות פירעון והוצאה לפועל</h4><span class="category-icon text-xs text-gray-500">▼</span></button>
                                <div id="category-bankruptcy" class="hidden">
                                    <a href="bankruptcy.html" class="block px-4 py-2 pr-8 text-sm text-gray-700 hover:bg-[#c9a55c] hover:text-[#1a1a1a] text-right">חדלות פירעון / פשיטת רגל</a>
                                    <a href="execution.html" class="block px-4 py-2 pr-8 text-sm text-gray-700 hover:bg-[#c9a55c] hover:text-[#1a1a1a] text-right">הוצאה לפועל</a>
                                    <a href="debt-cancellation.html" class="block px-4 py-2 pr-8 text-sm text-gray-700 hover:bg-[#c9a55c] hover:text-[#1a1a1a] text-right">מחיקת חובות</a>
                                    <a href="discharge.html" class="block px-4 py-2 pr-8 text-sm text-gray-700 hover:bg-[#c9a55c] hover:text-[#1a1a1a] text-right">קבלת הפטר</a>
                                </div>
                            </div>
                            <div class="mt-2"><a href="practice-areas.html#civil-law" class="block w-full px-4 py-2 text-right hover:bg-gray-50"><h4 class="font-bold text-[#1a1a1a] text-sm">דיני חוזים ודיני מקרקעין</h4></a></div>
                        </div>
                    </li>
                    <li><a href="articles.html" title="פרסומים מקצועיים" class="inline-flex items-center justify-center rounded-full border border-[#c9a55c] bg-[#1a1a1a] px-4 py-2 text-sm font-medium text-[#c9a55c] shadow-sm transition-all duration-300 hover:bg-[#c9a55c] hover:text-[#1a1a1a]">פרסומים מקצועיים</a></li>
                    <li><a href="blog.html" title="בלוג" class="inline-flex items-center justify-center rounded-full border border-[#c9a55c] bg-[#1a1a1a] px-4 py-2 text-sm font-medium text-[#c9a55c] shadow-sm transition-all duration-300 hover:bg-[#c9a55c] hover:text-[#1a1a1a]">בלוג</a></li>
                    <li><a href="https://www.lawreviews.co.il/provider/gil-yaniv#reviews" title="ביקורות מלקוחות" class="inline-flex items-center justify-center rounded-full border border-[#c9a55c] bg-[#1a1a1a] px-4 py-2 text-sm font-medium text-[#c9a55c] shadow-sm transition-all duration-300 hover:bg-[#c9a55c] hover:text-[#1a1a1a]" target="_blank" rel="noopener">ביקורות מלקוחות</a></li>
                    <li><a href="contact.html" title="צור קשר" class="inline-flex items-center justify-center rounded-full border border-[#c9a55c] bg-[#1a1a1a] px-4 py-2 text-sm font-medium text-[#c9a55c] shadow-sm transition-all duration-300 hover:bg-[#c9a55c] hover:text-[#1a1a1a]">צור קשר</a></li>
                </ul>
            </nav>
            <button id="mobile-menu-button" class="md:hidden flex items-center p-3 rounded-md text-gray-600 hover:text-blue-600 min-w-[48px] min-h-[48px]" aria-label="פתח תפריט נייד">
                <svg xmlns="http://www.w3.org/2000/svg" class="h-6 w-6" fill="none" viewBox="0 0 24 24" stroke="currentColor"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M4 6h16M4 12h16M4 18h16" /></svg>
            </button>
            <div class="flex items-center space-x-4 space-x-reverse">
                <a class="inline-flex items-center justify-center rounded-full border border-[#b8941f] bg-[#1a1a1a] px-5 py-3 text-base font-medium text-[#b8941f] shadow-sm transition-all duration-300 hover:bg-[#b8941f] hover:text-[#1a1a1a]" href="tel:0548184581" title="התקשר 054-8184581">התקשר עכשיו</a>
            </div>
        </div>
    </header>

    <!-- Mobile Menu - matches index.html -->
    <div id="mobile-menu" class="mobile-menu-overlay">
        <div class="flex justify-between items-center p-6 border-b">
            <div class="flex items-center gap-2">
                <a href="index.html" title="דף הבית"><img src="images/new-logo.png" alt="לוגו משרד עו״ד יניב גיל" class="h-10 w-auto"></a>
                <a href="index.html" title="דף הבית"><img src="images/Yaniv-Gil-Law-Office-Notary_text.avif" alt="משרד עו״ד יניב גיל ושות׳" class="h-8 w-auto"></a>
            </div>
            <button id="close-mobile-menu" class="text-gray-500 hover:text-gray-700 min-w-[48px] min-h-[48px] flex items-center justify-center" aria-label="סגור תפריט נייד">
                <svg xmlns="http://www.w3.org/2000/svg" class="h-6 w-6" fill="none" viewBox="0 0 24 24" stroke="currentColor"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M6 18L18 6M6 6l12 12" /></svg>
            </button>
        </div>
        <nav class="p-6">
            <ul class="space-y-4">
                <li><a href="index.html" title="דף הבית" class="block text-lg text-gray-800 hover:text-[#b8941f] mobile-menu-link">בית</a></li>
                <li><a href="index.html#about" title="אודות המשרד" class="block text-lg text-gray-800 hover:text-[#b8941f] mobile-menu-link">אודות המשרד</a></li>
                <li><a href="attorneys.html" title="עורכי דין" class="block text-lg text-gray-800 hover:text-[#b8941f] mobile-menu-link">עורכי דין</a></li>
                <li>
                    <button onclick="toggleNavDropdown('mobile-practice-dropdown')" class="mobile-category-btn block w-full text-right text-lg">תחומי התמחות <span class="mobile-category-chevron">▼</span></button>
                    <div id="mobile-practice-dropdown" class="hidden w-full mt-2 mobile-practice-container">
                        <div class="mb-3">
                            <button onclick="toggleCategoryDropdown('mobile-category-family', event)" class="mobile-category-btn"><span>דיני משפחה</span><span class="mobile-category-chevron">▼</span></button>
                            <div id="mobile-category-family" class="hidden mt-2 pr-4">
                                <a href="divorce-guide.html" class="block py-2 px-4 text-sm text-yellow-400 font-bold hover:text-[#c9a55c]">המדריך המלא לגירושין</a>
                                <a href="heskem-gerushin.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">הסכם גירושין</a>
                                <a href="gerushin-beit-din-rabani.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">בית דין רבני</a>
                                <a href="mishmoret-yeladim.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">משמורת ילדים</a>
                                <a href="mezonot-yeladim.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">מזונות ילדים</a>
                                <a href="mezonot-isha.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">מזונות אישה</a>
                                <a href="chalokat-rechush-gerushin.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">חלוקת רכוש</a>
                                <a href="heskem-mamon.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">הסכם ממון</a>
                                <a href="sarvanut-get-agunot.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">סרבנות גט</a>
                                <a href="yaduyim-batzibur.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">ידועים בציבור</a>
                                <a href="divorce.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">גירושין</a>
                            </div>
                        </div>
                        <div class="mb-3">
                            <button onclick="toggleCategoryDropdown('mobile-category-inheritance', event)" class="mobile-category-btn"><span>ירושות וצוואות</span><span class="mobile-category-chevron">▼</span></button>
                            <div id="mobile-category-inheritance" class="hidden mt-2 pr-4">
                                <a href="will-writing.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">עריכת צוואה</a>
                                <a href="inheritance-order.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">צו ירושה</a>
                                <a href="will-probate.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">צו קיום צוואה</a>
                            </div>
                        </div>
                        <div class="mb-3">
                            <button onclick="toggleCategoryDropdown('mobile-category-bankruptcy', event)" class="mobile-category-btn"><span>חדלות פירעון והוצאה לפועל</span><span class="mobile-category-chevron">▼</span></button>
                            <div id="mobile-category-bankruptcy" class="hidden mt-2 pr-4">
                                <a href="bankruptcy.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">חדלות פירעון / פשיטת רגל</a>
                                <a href="execution.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">הוצאה לפועל</a>
                                <a href="debt-cancellation.html" class="block py-2 px-4 text-sm text-gray-300 hover:text-[#c9a55c]">מחיקת חובות</a>
                            </div>
                        </div>
                        <div class="mb-3"><a href="practice-areas.html#civil-law" class="mobile-category-btn block w-full text-right">דיני חוזים ודיני מקרקעין</a></div>
                    </div>
                </li>
                <li><a href="articles.html" title="פרסומים מקצועיים" class="block text-lg text-gray-800 hover:text-[#b8941f] mobile-menu-link">פרסומים מקצועיים</a></li>
                <li><a href="blog.html" title="בלוג" class="block text-lg text-gray-800 hover:text-[#b8941f] mobile-menu-link">בלוג</a></li>
                <li><a href="contact.html" title="צור קשר" class="block text-lg text-gray-800 hover:text-[#b8941f] mobile-menu-link">צור קשר</a></li>
            </ul>
            <div class="mt-8"><a class="block w-full text-center py-3 px-4 bg-[#1a1a1a] text-[#b8941f] border border-[#b8941f] rounded-full hover:bg-[#b8941f] hover:text-[#1a1a1a]" href="tel:0548184581" title="התקשר 054-8184581">התקשר עכשיו</a></div>
        </nav>
    </div>

    <!-- Main Content -->
    <main class="pt-24 pb-16">
        <div class="container mx-auto px-4 max-w-4xl">
            <!-- Breadcrumbs -->
            {breadcrumb_html}

            <!-- Back Link -->
            <a href="divorce-guide.html" class="back-link">
                <svg class="w-4 h-4 transform rotate-180" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9 5l7 7-7 7"/></svg>
                חזרה למדריך הגירושין
            </a>

            <!-- Article Content -->
            <article class="bg-white rounded-xl shadow-lg p-6 md:p-8 content-section">
                {content}
            </article>

            <!-- Sub-pages Navigation -->
            {sub_pages_html}

            <!-- CTA Section -->
            <div class="cta-box mt-8">
                <h3 class="text-xl font-bold mb-4">צריכים ייעוץ משפטי?</h3>
                <p class="mb-4">עו"ד יניב גיל מתמחה בדיני משפחה ויעניק לכם ליווי מקצועי ואישי</p>
                <div class="flex flex-col sm:flex-row gap-4 justify-center">
                    <a href="tel:0548184581" class="cta-btn-call inline-flex items-center justify-center gap-2 px-6 py-3 rounded-full font-bold transition-colors">
                        <svg class="w-5 h-5" fill="currentColor" viewBox="0 0 24 24"><path d="M6.62 10.79c1.44 2.83 3.76 5.14 6.59 6.59l2.2-2.2c.27-.27.67-.36 1.02-.24 1.12.37 2.33.57 3.57.57.55 0 1 .45 1 1V20c0 .55-.45 1-1 1-9.39 0-17-7.61-17-17 0-.55.45-1 1-1h3.5c.55 0 1 .45 1 1 0 1.25.2 2.45.57 3.57.11.35.03.74-.25 1.02l-2.2 2.2z"/></svg>
                        התקשרו עכשיו
                    </a>
                    <a href="https://wa.me/972548184581" target="_blank" rel="noopener" class="cta-btn-whatsapp inline-flex items-center justify-center gap-2 px-6 py-3 rounded-full font-bold transition-colors">
                        <svg class="w-5 h-5" fill="currentColor" viewBox="0 0 24 24"><path d="M17.472 14.382c-.297-.149-1.758-.867-2.03-.967-.273-.099-.471-.148-.67.15-.197.297-.767.966-.94 1.164-.173.199-.347.223-.644.075-.297-.15-1.255-.463-2.39-1.475-.883-.788-1.48-1.761-1.653-2.059-.173-.297-.018-.458.13-.606.134-.133.298-.347.446-.52.149-.174.198-.298.298-.497.099-.198.05-.371-.025-.52-.075-.149-.669-1.612-.916-2.207-.242-.579-.487-.5-.669-.51-.173-.008-.371-.01-.57-.01-.198 0-.52.074-.792.372-.272.297-1.04 1.016-1.04 2.479 0 1.462 1.065 2.875 1.213 3.074.149.198 2.096 3.2 5.077 4.487.709.306 1.262.489 1.694.625.712.227 1.36.195 1.871.118.571-.085 1.758-.719 2.006-1.413.248-.694.248-1.289.173-1.413-.074-.124-.272-.198-.57-.347z"/></svg>
                        WhatsApp
                    </a>
                </div>
            </div>
        </div>
    </main>

    <!-- Footer - matches index.html exactly -->
    <footer class="bg-gray-900 text-gray-300 py-8" role="contentinfo">
        <div class="container mx-auto px-4">
            <div class="grid grid-cols-1 md:grid-cols-4 gap-8 text-sm">
                <div>
                    <h3 class="text-lg font-bold text-white mb-4">המשרד</h3>
                    <ul class="space-y-2">
                        <li><a href="index.html" title="דף הבית" class="hover:text-[#b8941f] transition-colors duration-200">בית</a></li>
                        <li><a href="index.html#about" title="אודות המשרד" class="hover:text-[#b8941f] transition-colors duration-200">אודות המשרד</a></li>
                        <li><a href="attorneys.html" title="עורכי דין" class="hover:text-[#b8941f] transition-colors duration-200">עורכי דין</a></li>
                        <li><a href="practice-areas.html" title="תחומי התמחות" class="hover:text-[#b8941f] transition-colors duration-200">תחומי התמחות</a></li>
                        <li><a href="articles.html" title="פרסומים מקצועיים" class="hover:text-[#b8941f] transition-colors duration-200">פרסומים מקצועיים</a></li>
                        <li><a href="blog.html" title="בלוג" class="hover:text-[#b8941f] transition-colors duration-200">בלוג</a></li>
                        <li><a href="https://www.lawreviews.co.il/provider/gil-yaniv#reviews" title="ביקורות מלקוחות" class="hover:text-[#b8941f] transition-colors duration-200" target="_blank" rel="noopener">ביקורות מלקוחות</a></li>
                        <li><a href="contact.html" title="צור קשר" class="hover:text-[#b8941f] transition-colors duration-200">צור קשר</a></li>
                    </ul>
                </div>
                <div>
                    <h3 class="text-lg font-bold text-white mb-4">תחומי התמחות</h3>
                    <ul class="space-y-2">
                        <li><a href="practice-areas.html#family-law" title="דיני משפחה" class="hover:text-[#b8941f] transition-colors duration-200">דיני משפחה, ירושות וצוואות</a></li>
                        <li><a href="practice-areas.html#bankruptcy" title="חדלות פירעון" class="hover:text-[#b8941f] transition-colors duration-200">חדלות פירעון והוצאה לפועל</a></li>
                        <li><a href="practice-areas.html#civil-law" title="סדר דין אזרחי" class="hover:text-[#b8941f] transition-colors duration-200">סדר דין אזרחי ומקרקעין</a></li>
                    </ul>
                </div>
                <div>
                    <h3 class="text-lg font-bold text-white mb-4">יצירת קשר</h3>
                    <div class="space-y-4">
                        <p><strong>טלפון ראשי:</strong><br><a href="tel:03-6092414" title="התקשר 03-6092414" class="hover:text-[#b8941f]">03-6092414</a></p>
                        <p><strong>נייד:</strong><br><a href="tel:054-8184581" title="התקשר 054-8184581" class="hover:text-[#b8941f]">054-8184581</a></p>
                        <p><strong>פקס:</strong><br><a href="tel:03-6092413" title="פקס 03-6092413" class="hover:text-[#b8941f]">03-6092413</a></p>
                        <p><strong>אימייל:</strong><br><a href="mailto:yanivgil@gmail.com" class="hover:text-[#b8941f]">yanivgil@gmail.com</a></p>
                        <p><strong>כתובת:</strong><br>מגדל "WE-TLV"<br>דרך מנחם בגין 150<br>קומה 7, תל-אביב</p>
                    </div>
                </div>
            </div>
            <div class="mt-8 pt-6 border-t border-gray-700 flex flex-col md:flex-row justify-between items-center text-xs text-gray-500">
                <div class="flex flex-col items-center md:items-start">
                    <p>&copy; 2026 משרד עו״ד יניב גיל. כל הזכויות שמורות.</p>
                    <p class="mt-1">האתר נבנה ע"י <a href="https://eladdigital.vercel.app/" target="_blank" rel="noopener" class="text-gray-300 hover:text-[#b8941f] transition-colors underline">eladdigital</a></p>
                </div>
                <div class="flex gap-4 mt-4 md:mt-0">
                    <a href="privacy-policy.html" class="hover:text-[#b8941f] transition-colors duration-200">מדיניות פרטיות</a>
                    <a href="terms-of-use.html" class="hover:text-[#b8941f] transition-colors duration-200">תנאי שימוש</a>
                    <a href="accessibility.html" class="hover:text-[#b8941f] transition-colors duration-200">הצהרת נגישות</a>
                </div>
            </div>
        </div>
    </footer>

    <script>
        // Header scroll effect
        window.addEventListener('scroll', function() {{
            const header = document.getElementById('site-header');
            if (window.scrollY > 50) {{ header.classList.add('py-3', 'shadow-xl', 'scrolled'); header.classList.remove('py-6', 'shadow-md'); }}
            else {{ header.classList.remove('py-3', 'shadow-xl', 'scrolled'); header.classList.add('py-6', 'shadow-md'); }}
        }});
        // Mobile menu toggle
        document.addEventListener('DOMContentLoaded', function() {{
            const mobileMenuButton = document.getElementById('mobile-menu-button');
            const closeMenuButton = document.getElementById('close-mobile-menu');
            const mobileMenu = document.getElementById('mobile-menu');
            if (mobileMenuButton) mobileMenuButton.addEventListener('click', function() {{ mobileMenu.classList.add('active'); document.body.style.overflow = 'hidden'; }});
            if (closeMenuButton) closeMenuButton.addEventListener('click', function() {{ mobileMenu.classList.remove('active'); document.body.style.overflow = ''; }});
        }});
        // Dropdown toggles
        function toggleNavDropdown(dropdownId) {{
            const dropdown = document.getElementById(dropdownId);
            if (!dropdown) return;
            document.querySelectorAll('[id$="-dropdown"]').forEach(dd => {{ if (dd.id !== dropdownId) dd.classList.add('hidden'); }});
            dropdown.classList.toggle('hidden');
        }}
        function toggleCategoryDropdown(categoryId, event) {{
            if (event) event.stopPropagation();
            const category = document.getElementById(categoryId);
            if (!category) return;
            const button = event ? event.target.closest('button') : null;
            const icon = button ? button.querySelector('.category-icon') : null;
            if (category.classList.contains('hidden')) {{ category.classList.remove('hidden'); if (icon) icon.classList.add('open'); }}
            else {{ category.classList.add('hidden'); if (icon) icon.classList.remove('open'); }}
        }}
        function toggleSubCategory(subCategoryId, event) {{
            if (event) event.stopPropagation();
            const subCategory = document.getElementById(subCategoryId);
            if (!subCategory) return;
            const button = event ? event.target.closest('button') : null;
            if (subCategory.classList.contains('hidden')) {{ subCategory.classList.remove('hidden'); if (button) button.textContent = '▲'; }}
            else {{ subCategory.classList.add('hidden'); if (button) button.textContent = '▼'; }}
        }}
        // Close dropdown when clicking outside
        document.addEventListener('click', function(e) {{
            if (!e.target.closest('.relative.group') && !e.target.closest('[id$="-dropdown"]')) {{
                document.querySelectorAll('[id$="-dropdown"]').forEach(dd => dd.classList.add('hidden'));
            }}
        }});
    </script>

    <!-- Floating Buttons -->
    <a href="tel:0548184581" title="התקשר עכשיו" class="call-float" aria-label="התקשר 054-8184581">
        <svg fill="currentColor" viewBox="0 0 24 24" class="w-5 h-5"><path d="M6.62 10.79c1.44 2.83 3.76 5.14 6.59 6.59l2.2-2.2c.27-.27.67-.36 1.02-.24 1.12.37 2.33.57 3.57.57.55 0 1 .45 1 1V20c0 .55-.45 1-1 1-9.39 0-17-7.61-17-17 0-.55.45-1 1-1h3.5c.55 0 1 .45 1 1 0 1.25.2 2.45.57 3.57.11.35.03.74-.25 1.02l-2.2 2.2z"/></svg>
    </a>
    <a href="https://wa.me/972548184581" title="שלח WhatsApp" class="whatsapp-float" target="_blank" rel="noopener" aria-label="שלח WhatsApp">
        <svg fill="currentColor" viewBox="0 0 24 24" class="w-6 h-6"><path d="M17.472 14.382c-.297-.149-1.758-.867-2.03-.967-.273-.099-.471-.148-.67.15-.197.297-.767.966-.94 1.164-.173.199-.347.223-.644.075-.297-.15-1.255-.463-2.39-1.475-.883-.788-1.48-1.761-1.653-2.059-.173-.297-.018-.458.13-.606.134-.133.298-.347.446-.52.149-.174.198-.298.298-.497.099-.198.05-.371-.025-.52-.075-.149-.669-1.612-.916-2.207-.242-.579-.487-.5-.669-.51-.173-.008-.371-.01-.57-.01-.198 0-.52.074-.792.372-.272.297-1.04 1.016-1.04 2.479 0 1.462 1.065 2.875 1.213 3.074.149.198 2.096 3.2 5.077 4.487.709.306 1.262.489 1.694.625.712.227 1.36.195 1.871.118.571-.085 1.758-.719 2.006-1.413.248-.694.248-1.289.173-1.413-.074-.124-.272-.198-.57-.347z"/></svg>
    </a>
    <button class="back-to-top" aria-label="חזרה לראש הדף" title="חזרה למעלה">
        <svg fill="none" stroke="currentColor" viewBox="0 0 24 24" class="w-6 h-6"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M5 10l7-7m0 0l7 7m-7-7v18"/></svg>
    </button>

    <!-- Mobile CTA Bar -->
    <div class="mobile-cta-bar">
        <a href="tel:0548184581" class="cta-call" aria-label="התקשר עכשיו">
            <svg viewBox="0 0 24 24"><path d="M20.01 15.38c-1.23 0-2.42-.2-3.53-.56a.977.977 0 00-1.01.24l-1.57 1.97c-2.83-1.35-5.48-3.9-6.89-6.83l1.95-1.66c.27-.28.35-.67.24-1.02-.37-1.11-.56-2.3-.56-3.53 0-.54-.45-.99-.99-.99H4.19C3.65 3 3 3.24 3 3.99 3 13.28 10.73 21 20.01 21c.71 0 .99-.63.99-1.18v-3.45c0-.54-.45-.99-.99-.99z"/></svg>
            התקשר עכשיו
        </a>
        <a href="https://wa.me/972548184581" class="cta-whatsapp" target="_blank" rel="noopener" aria-label="WhatsApp">
            <svg viewBox="0 0 24 24"><path d="M17.472 14.382c-.297-.149-1.758-.867-2.03-.967-.273-.099-.471-.148-.67.15-.197.297-.767.966-.94 1.164-.173.199-.347.223-.644.075-.297-.15-1.255-.463-2.39-1.475-.883-.788-1.48-1.761-1.653-2.059-.173-.297-.018-.458.13-.606.134-.133.298-.347.446-.52.149-.174.198-.298.298-.497.099-.198.05-.371-.025-.52-.075-.149-.669-1.612-.916-2.207-.242-.579-.487-.5-.669-.51-.173-.008-.371-.01-.57-.01-.198 0-.52.074-.792.372-.272.297-1.04 1.016-1.04 2.479 0 1.462 1.065 2.875 1.213 3.074.149.198 2.096 3.2 5.077 4.487.709.306 1.262.489 1.694.625.712.227 1.36.195 1.871.118.571-.085 1.758-.719 2.006-1.413.248-.694.248-1.289.173-1.413-.074-.124-.272-.198-.57-.347z"/></svg>
            WhatsApp
        </a>
    </div>

    <script src="js/back-to-top.js"></script>
    <script src="js/progress-bar.js"></script>
    <script src="js/conversion-tracking.js" defer></script>
</body>
</html>'''


def get_element_index(doc, element):
    """Get the index of an element in the document body."""
    body = doc.element.body
    for i, child in enumerate(body):
        if child is element:
            return i
    return -1


def iter_block_items(doc):
    """
    Iterate through all block-level elements (paragraphs and tables) in document order.
    This ensures tables are processed in their correct position within the content.
    """
    parent = doc.element.body

    for child in parent.iterchildren():
        if child.tag == qn('w:p'):
            # It's a paragraph
            for para in doc.paragraphs:
                if para._element is child:
                    yield ('paragraph', para)
                    break
        elif child.tag == qn('w:tbl'):
            # It's a table
            for table in doc.tables:
                if table._element is child:
                    yield ('table', table)
                    break


def convert_paragraph_to_html(text, style_name='', is_bold=False):
    """Convert a single paragraph to HTML. Returns (html, element_type)."""
    text = text.strip()
    if not text:
        return '', 'empty'

    # Handle headings
    if style_name == 'Heading 1':
        return f'<h1 class="text-3xl md:text-4xl font-bold text-[#1a1a1a] mb-6">{text}</h1>\n', 'h1'
    elif style_name == 'Heading 2':
        # Create anchor ID from text
        anchor_id = re.sub(r'[^\w\s\u0590-\u05FF]', '', text).replace(' ', '-')[:50]
        return f'<h2 id="{anchor_id}" class="text-2xl font-bold text-[#1a1a1a] mt-8 mb-4 pb-2 border-b-2 border-[#c9a55c]">{text}</h2>\n', 'h2'
    elif style_name == 'Heading 3':
        return f'<h3 class="text-xl font-semibold text-gray-700 mt-6 mb-3">{text}</h3>\n', 'h3'

    # Handle bullet points
    if text.startswith('•') or text.startswith('-') or text.startswith('–'):
        item_text = text.lstrip('•-–').strip()
        return f'<li>{item_text}</li>\n', 'bullet'

    # Handle numbered lists
    if re.match(r'^\d+[\.\)]\s', text):
        item_text = re.sub(r'^\d+[\.\)]\s*', '', text)
        return f'<li>{item_text}</li>\n', 'number'

    # Bold text becomes a subheading if it's standalone
    if is_bold and len(text) < 100:
        return f'<h4 class="text-lg font-semibold text-gray-800 mt-4 mb-2">{text}</h4>\n', 'h4'

    # Regular paragraph
    return f'<p class="text-gray-600 leading-relaxed mb-4">{text}</p>\n', 'paragraph'


def convert_table_to_html(table):
    """Convert a Word table to HTML, handling various table types."""
    rows = []
    cell_texts = []

    for row in table.rows:
        for cell in row.cells:
            cell_texts.append(cell.text.strip())

    full_text = ' '.join(cell_texts)

    # Check if it's a single-cell table (callout box, CTA, or author box)
    if len(table.rows) == 1 and len(table.columns) == 1:
        cell_text = table.rows[0].cells[0].text.strip()

        # Format the text with line breaks
        formatted_text = cell_text.replace('\n', '<br>')

        # CTA box (phone, consultation)
        if any(keyword in cell_text for keyword in ['📞', 'לייעוץ', 'התקשרו', 'יצירת קשר', 'לפגישה']):
            return f'''<div class="cta-box my-6">
                <div class="text-lg">{formatted_text}</div>
            </div>\n'''

        # Author box
        elif any(keyword in cell_text[:50] for keyword in ['מאמר זה נכתב', 'נכתב על ידי', 'עו"ד יניב', 'עורך דין']):
            return f'''<div class="author-box">
                <div class="text-gray-600">{formatted_text}</div>
            </div>\n'''

        # Tip/Important box
        elif any(keyword in cell_text for keyword in ['חשוב', 'שימו לב', 'טיפ', 'הערה']):
            return f'''<div class="tip-box">
                <div class="text-gray-700">{formatted_text}</div>
            </div>\n'''

        # Warning box
        elif any(keyword in cell_text for keyword in ['אזהרה', 'זהירות', 'סיכון']):
            return f'''<div class="warning-box">
                <div class="text-gray-700">{formatted_text}</div>
            </div>\n'''

        # General info box
        else:
            return f'''<div class="info-box">
                <div class="text-gray-700">{formatted_text}</div>
            </div>\n'''

    # Regular multi-cell table
    for i, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace('\n', '<br>')
            if i == 0:  # Header row
                cells.append(f'<th>{cell_text}</th>')
            else:
                cells.append(f'<td>{cell_text}</td>')

        if i == 0:
            rows.append('<thead><tr>' + ''.join(cells) + '</tr></thead>')
        else:
            if i == 1:
                rows.append('<tbody>')
            rows.append('<tr>' + ''.join(cells) + '</tr>')

    if len(table.rows) > 1:
        rows.append('</tbody>')

    return f'<table class="content-table">\n{"".join(rows)}\n</table>\n'


def is_page_separator(text):
    """Check if text is a URL/page separator line."""
    # Match lines that contain URL: followed by a path
    return 'URL:' in text and '/' in text


def is_metadata_line(text):
    """Check if text is a metadata line (not content)."""
    # Only skip truly empty or separator lines
    if not text.strip():
        return True
    if text.strip() in ['─────────────────────────────────────────', '~', '───']:
        return True
    if text.startswith('Pillar מרכזי:'):
        return True
    # Filter out "עמוד:" lines that indicate page type (תת-עמוד, Pillar ראשי, etc.)
    if text.strip().startswith('עמוד:'):
        return True
    return False


def parse_document_full(doc):
    """
    Parse a Word document and extract ALL content for each page.
    Preserves document order including tables interspersed with paragraphs.
    """
    pages = []
    current_page = None
    current_content = []
    in_list = False
    list_type = None  # 'ul' or 'ol'

    # Iterate through all block elements in order
    for item_type, item in iter_block_items(doc):

        if item_type == 'paragraph':
            para = item
            text = para.text.strip()
            style_name = para.style.name if para.style else ''

            # Check if paragraph is bold
            is_bold = False
            if para.runs:
                is_bold = all(run.bold for run in para.runs if run.text.strip())

            # Check for page marker (URL line)
            if is_page_separator(text):
                # Save previous page if exists
                if current_page:
                    if in_list:
                        current_content.append(f'</{list_type}>\n')
                        in_list = False
                    current_page['content'] = ''.join(current_content)
                    pages.append(current_page)

                # Parse URL and keyword
                # Format: עמוד: Pillar ראשי  |  URL: /gerushin-beit-din-rabani.html  |  מילת מפתח: גירושין בית דין רבני
                url_match = re.search(r'URL:\s*/?([^\s|]+)', text)
                keyword_match = re.search(r'מילת מפתח:\s*(.+?)(?:\s*$|\|)', text)

                url = ''
                if url_match:
                    url = url_match.group(1).strip()
                    # Remove leading slash if present
                    url = url.lstrip('/')
                    if not url.endswith('.html'):
                        url = url + '.html'

                current_page = {
                    'url': url,
                    'keyword': keyword_match.group(1).strip() if keyword_match else '',
                    'title': '',
                    'description': ''
                }
                current_content = []
                in_list = False
                continue

            # Skip if no current page started
            if not current_page:
                continue

            # Skip metadata lines
            if is_metadata_line(text):
                continue

            # Convert paragraph
            html, element_type = convert_paragraph_to_html(text, style_name, is_bold)

            if not html or element_type == 'empty':
                continue

            # Handle list wrapping
            is_list_item = element_type in ['bullet', 'number']

            if is_list_item:
                new_list_type = 'ul' if element_type == 'bullet' else 'ol'
                if not in_list:
                    list_type = new_list_type
                    current_content.append(f'<{list_type} class="space-y-2 mb-4">\n')
                    in_list = True
                elif list_type != new_list_type:
                    # Close old list, start new
                    current_content.append(f'</{list_type}>\n')
                    list_type = new_list_type
                    current_content.append(f'<{list_type} class="space-y-2 mb-4">\n')
            elif in_list:
                current_content.append(f'</{list_type}>\n')
                in_list = False

            # Extract title from first H1
            if element_type == 'h1' and not current_page['title']:
                current_page['title'] = text
                # Generate description from title
                current_page['description'] = f'{text} - מדריך מקיף מאת עו"ד יניב גיל, מומחה לדיני משפחה בתל אביב. מידע משפטי מקצועי וטיפים חשובים.'

            current_content.append(html)

        elif item_type == 'table':
            table = item

            # If no page started yet, skip
            if not current_page:
                continue

            # Close any open list
            if in_list:
                current_content.append(f'</{list_type}>\n')
                in_list = False

            # Convert and add table
            table_html = convert_table_to_html(table)
            current_content.append(table_html)

    # Save last page
    if current_page:
        if in_list:
            current_content.append(f'</{list_type}>\n')
        current_page['content'] = ''.join(current_content)
        pages.append(current_page)

    return pages


def process_document(filepath, pillar_name, pillar_url):
    """Process a single Word document and generate HTML pages."""
    print(f'\nProcessing: {filepath}')

    doc = Document(filepath)
    pages = parse_document_full(doc)

    print(f'  Found {len(pages)} pages')

    # Count total content
    total_paras = len(doc.paragraphs)
    total_tables = len(doc.tables)
    print(f'  Document has {total_paras} paragraphs and {total_tables} tables')

    created_files = []

    # Collect sub-pages for pillar page navigation
    sub_pages = []
    for i, page in enumerate(pages[1:], 1):  # Skip first (pillar) page
        if page.get('url') and page.get('title'):
            sub_pages.append({
                'url': page['url'],
                'title': page['title']
            })

    for i, page in enumerate(pages):
        if not page['url'] or not page['title']:
            print(f'  Skipping page {i+1}: missing URL or title')
            continue

        # Clean URL
        url = page['url'].strip('/')
        if not url.endswith('.html'):
            url += '.html'

        # Determine if this is a pillar or sub-page
        is_pillar = (i == 0)

        # Generate HTML content
        content = page['content']

        # Count content in this page
        content_length = len(content)
        h2_count = content.count('<h2')
        p_count = content.count('<p ')
        print(f'    Page {i+1}: {url} - {content_length} chars, {h2_count} h2s, {p_count} paragraphs')

        # Generate full HTML
        if is_pillar:
            html = get_html_template(
                title=page['title'],
                description=page['description'],
                keyword=page['keyword'],
                url=url,
                content=content,
                sub_pages=sub_pages  # Add sub-pages navigation to pillar
            )
        else:
            html = get_html_template(
                title=page['title'],
                description=page['description'],
                keyword=page['keyword'],
                url=url,
                content=content,
                breadcrumb_pillar=pillar_name,
                breadcrumb_pillar_url=pillar_url
            )

        # Write file
        output_path = os.path.join(OUTPUT_FOLDER, url)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)

        print(f'  Created: {url}')
        created_files.append({
            'url': url,
            'title': page['title'],
            'is_pillar': is_pillar
        })

    return created_files


def main():
    """Main function to process all Word documents."""
    print('=' * 60)
    print('Word to HTML Converter for yanivgil.co.il')
    print('FULL CONTENT VERSION - No shortening')
    print('=' * 60)

    # Get all pillar documents
    files = []
    for f in os.listdir(INPUT_FOLDER):
        if f.endswith('.docx') and 'Pillar' in f:
            # Skip duplicates and strategy doc
            if any(skip in f for skip in SKIP_FILES):
                continue
            files.append(f)

    files.sort()
    print(f'\nFound {len(files)} pillar documents to process')

    all_created_files = []
    pillar_info = {}  # Store pillar name -> URL mapping

    # First pass: identify pillar URLs
    for filename in files:
        filepath = os.path.join(INPUT_FOLDER, filename)
        doc = Document(filepath)

        pillar_name = filename.replace('Pillar ', '').replace(' COMPLETE.docx', '')

        # Find first URL (pillar page)
        for para in doc.paragraphs:
            text = para.text.strip()
            if is_page_separator(text):
                url_match = re.search(r'URL:\s*/?([^\s|]+)', text)
                if url_match:
                    pillar_url = url_match.group(1).strip().lstrip('/')
                    if not pillar_url.endswith('.html'):
                        pillar_url += '.html'
                    pillar_info[pillar_name] = pillar_url

                    # Also get Hebrew name from first H1
                    for p in doc.paragraphs:
                        if p.style and p.style.name == 'Heading 1' and p.text.strip():
                            pillar_info[pillar_name + '_hebrew'] = p.text.strip()
                            break
                    break

    # Second pass: process all documents
    for filename in files:
        filepath = os.path.join(INPUT_FOLDER, filename)
        pillar_name = filename.replace('Pillar ', '').replace(' COMPLETE.docx', '')
        pillar_url = pillar_info.get(pillar_name, '')
        pillar_hebrew = pillar_info.get(pillar_name + '_hebrew', pillar_name)

        created = process_document(filepath, pillar_hebrew, pillar_url)
        all_created_files.extend(created)

    print('\n' + '=' * 60)
    print(f'COMPLETE! Created {len(all_created_files)} HTML files')
    print('=' * 60)

    # Output list of URLs for sitemap
    print('\nURLs to add to sitemap.xml:')
    for file_info in all_created_files:
        priority = '0.9' if file_info['is_pillar'] else '0.7'
        print(f'  https://www.yanivgil.co.il/{file_info["url"]} (priority: {priority})')

    # Output pillar info for menu structure
    print('\n\nPillar pages for navigation menu:')
    for filename in files:
        pillar_name = filename.replace('Pillar ', '').replace(' COMPLETE.docx', '')
        pillar_url = pillar_info.get(pillar_name, '')
        pillar_hebrew = pillar_info.get(pillar_name + '_hebrew', pillar_name)
        print(f'  {pillar_hebrew}: {pillar_url}')

    return all_created_files


if __name__ == '__main__':
    main()
